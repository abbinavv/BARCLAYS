from flask import Flask, request, jsonify, send_file, render_template
from docx import Document
import pandas as pd
import PyPDF2
from docx import Document as DocxDocument
from pymongo import MongoClient
from bson import ObjectId
import tempfile
import os
from datetime import datetime
from jira import JIRA
from transformers import pipeline

app = Flask(__name__, static_folder='static', template_folder='templates')

# Connect to MongoDB for version control
client = MongoClient('mongodb://localhost:27017/')
db = client['recap_ai']
documents_collection = db['documents']

# JIRA configuration (replace with your credentials)
JIRA_SERVER = 'https://your-jira-instance.atlassian.net'
JIRA_USERNAME = 'your-username'
JIRA_API_TOKEN = 'your-api-token'
JIRA_PROJECT_KEY = 'YOUR_PROJECT_KEY'

# Initialize the NLP model for requirement classification
classifier = pipeline("text-classification", model="distilbert-base-uncased")

# Function to push user stories to JIRA
def push_to_jira(user_stories):
    try:
        jira = JIRA(server=JIRA_SERVER, basic_auth=(JIRA_USERNAME, JIRA_API_TOKEN))
        for story in user_stories:
            issue_dict = {
                'project': {'key': JIRA_PROJECT_KEY},
                'summary': story['User Story'],
                'description': story.get('Acceptance Criteria', 'TBD'),
                'issuetype': {'name': 'Story'},
                'priority': {'name': f"Priority {story.get('Priority', 0)}"}
            }
            jira.create_issue(fields=issue_dict)
        print("Successfully pushed user stories to JIRA.")
    except Exception as e:
        print(f"Error pushing to JIRA: {e}")

# Function to extract text from files (PDF, Word, Text)
def extract_text_from_file(file):
    try:
        if file.filename.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        elif file.filename.endswith('.docx'):
            doc = DocxDocument(file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        elif file.filename.endswith('.txt'):
            return file.read().decode('utf-8')
        else:
            return ""  # Unsupported file type
    except Exception as e:
        print(f"Error extracting text from file: {e}")
        return ""

# Function to process text and extract requirements (improved for resumes)
def process_text(text):
    # Split text into lines to identify sections
    lines = text.split('\n')
    functional = []
    non_functional = []
    clarifications = []
    priority_scores = []
    current_section = None

    # Keywords for identifying sections in a resume
    section_headers = {
        "objectives": ["objectives", "objective", "summary"],
        "skills": ["skills", "technical skills", "expertise"],
        "experience": ["experience", "work experience", "professional experience"],
        "education": ["education", "academic background"],
        "interests": ["interests", "hobbies"]
    }

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Identify section headers
        line_lower = line.lower()
        for section, keywords in section_headers.items():
            if any(keyword in line_lower for keyword in keywords):
                current_section = section
                continue

        # Process lines based on the current section
        if current_section == "objectives":
            # Treat objectives as high-level functional requirements
            functional.append(f"The system shall support the goal: {line}")
            priority_scores.append(("Must", 8))  # High priority for objectives
            # Check for vague terms
            if "innovative" in line_lower or "driven" in line_lower:
                clarifications.append(f"Please clarify: '{line}' - What specific goals or projects are intended?")

        elif current_section == "skills":
            # Treat skills as technical requirements (non-functional)
            non_functional.append(f"The system shall support the technology: {line}")
            priority_scores.append(("Should", 5))  # Medium priority for skills

        elif current_section == "experience":
            # Extract potential requirements from experience
            if "developed" in line_lower or "built" in line_lower:
                functional.append(f"The system shall include functionality similar to: {line}")
                priority_scores.append(("Could", 3))
            elif "managed" in line_lower or "organized" in line_lower:
                non_functional.append(f"The system shall support management tasks like: {line}")
                priority_scores.append(("Could", 3))

        elif current_section == "education":
            # Education might indicate domain knowledge
            non_functional.append(f"The system shall leverage knowledge from: {line}")
            priority_scores.append(("Could", 3))

        elif current_section == "interests":
            # Interests might indicate user preferences
            non_functional.append(f"The system may consider user interest: {line}")
            priority_scores.append(("Could", 3))

        else:
            # Fallback: use the NLP model for unclassified lines
            result = classifier(line)
            score = result[0]['score']
            if "shall" in line_lower or "must" in line_lower or score > 0.7:
                functional.append(line)
                priority_scores.append(("Must", 8))
            else:
                non_functional.append(line)
                priority_scores.append(("Could", 3))

            # Check for ambiguous terms
            if "fast" in line_lower or "secure" in line_lower:
                clarifications.append(f"Please clarify: '{line}' - What does 'fast' or 'secure' mean in this context?")

    # Aggregate priorities
    priority = {}
    for label, score in priority_scores:
        if label not in priority:
            priority[label] = score

    return {
        "functional": functional,
        "non_functional": non_functional,
        "priority": priority,
        "clarifications": clarifications
    }

# Function to generate a structured Word document (2-3 pages)
def generate_word_doc(requirements):
    doc = Document()
    
    doc.add_heading("Requirements Document", 0)
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This document outlines the requirements extracted using Recap AI, an AI-powered system for automated requirement gathering.")
    
    doc.add_heading("Functional Requirements", level=1)
    if requirements["functional"]:
        for req in requirements["functional"]:
            doc.add_paragraph(req, style='List Bullet')
    else:
        doc.add_paragraph("No functional requirements identified.")
    
    doc.add_heading("Non-Functional Requirements", level=1)
    if requirements["non_functional"]:
        for req in requirements["non_functional"]:
            doc.add_paragraph(req, style='List Bullet')
    else:
        doc.add_paragraph("No non-functional requirements identified.")
    
    doc.add_heading("Priority (MoSCoW Method)", level=1)
    for label, score in requirements["priority"].items():
        doc.add_paragraph(f"{label}: {score}")

    if requirements["clarifications"]:
        doc.add_heading("Clarifications Needed", level=1)
        for clarification in requirements["clarifications"]:
            doc.add_paragraph(clarification, style='List Bullet')

    filename = f"requirements_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)

    with open(filename, 'rb') as f:
        file_data = f.read()
    documents_collection.insert_one({
        "type": "word",
        "filename": filename,
        "data": file_data,
        "timestamp": datetime.now(),
        "version": documents_collection.count_documents({"type": "word"}) + 1
    })

    return filename

# Function to generate Excel user stories for JIRA backlog
def generate_excel_user_stories(requirements):
    data = []
    for req in requirements["functional"]:
        user_story = {
            "User Story": f"As a user, I want {req} so that I can achieve my goal.",
            "Acceptance Criteria": "TBD",
            "Priority": requirements["priority"].get("Must", 0)
        }
        data.append(user_story)
    
    push_to_jira(data)

    df = pd.DataFrame(data)
    filename = f"user_stories_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    df.to_excel(filename, index=False)

    with open(filename, 'rb') as f:
        file_data = f.read()
    documents_collection.insert_one({
        "type": "excel",
        "filename": filename,
        "data": file_data,
        "timestamp": datetime.now(),
        "version": documents_collection.count_documents({"type": "excel"}) + 1
    })

    return filename

# Homepage route
@app.route('/')
def index():
    return render_template('index.html')

# Analyze page route
@app.route('/analyze')
def analyze_page():
    return render_template('analyze.html')

# Process combined text and file input
@app.route('/process-input', methods=['POST'])
def process_input():
    text = request.form.get('text', '')
    file = request.files.get('file')

    combined_text = text

    if file and file.filename:
        file_text = extract_text_from_file(file)
        if file_text:
            combined_text += "\n" + file_text

    if not combined_text.strip():
        return jsonify({"error": "No input provided"}), 400

    requirements = process_text(combined_text)
    return jsonify({"requirements": requirements})

# Download Word document
@app.route('/download-word', methods=['POST'])
def download_word():
    data = request.json
    requirements = data.get('requirements')
    filename = generate_word_doc(requirements)
    return send_file(filename, as_attachment=True)

# Download Excel file
@app.route('/download-excel', methods=['POST'])
def download_excel():
    data = request.json
    requirements = data.get('requirements')
    filename = generate_excel_user_stories(requirements)
    return send_file(filename, as_attachment=True)

# View document versions
@app.route('/versions')
def view_versions():
    versions = list(documents_collection.find().sort("timestamp", -1))
    return render_template('versions.html', versions=versions)

# Download a specific version of a document
@app.route('/download-version/<version_id>')
def download_version(version_id):
    version = documents_collection.find_one({"_id": ObjectId(version_id)})
    if not version:
        return "Version not found", 404
    with tempfile.NamedTemporaryFile(delete=False, suffix='.' + version['type']) as temp_file:
        temp_file.write(version['data'])
        temp_file_path = temp_file.name
    return send_file(temp_file_path, as_attachment=True, download_name=version['filename'])

if __name__ == '__main__':
    app.run(debug=True)